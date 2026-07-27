from __future__ import annotations

from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.extractors.base import BaseExtractor, CodeBlock, ExtractedDocument, Heading, ImageMeta, Paragraph, Table


class DocxExtractor(BaseExtractor):
    HEADING_MAP = {
        WD_ALIGN_PARAGRAPH.CENTER: None,
    }

    def extract(self, file_bytes: bytes, filename: str, content_type: str) -> ExtractedDocument:
        doc_hash = self._hash_bytes(file_bytes)
        headings: list[Heading] = []
        paragraphs: list[Paragraph] = []
        tables: list[Table] = []
        images: list[ImageMeta] = []
        code_blocks: list[CodeBlock] = []
        metadata: dict[str, Any] = {}
        title = ""
        author = ""

        try:
            with DocxDocument(reader=file_bytes) as doc:
                core = doc.core_properties
                if core.title:
                    title = core.title
                if core.author:
                    author = core.author
                if core.subject:
                    metadata["subject"] = core.subject
                if core.keywords:
                    metadata["keywords"] = core.keywords
                if core.category:
                    metadata["category"] = core.category
                if core.comments:
                    metadata["comments"] = core.comments

                heading_style_map = {
                    "heading 1": 1,
                    "heading 2": 2,
                    "heading 3": 3,
                    "heading 4": 4,
                    "heading 5": 5,
                    "heading 6": 6,
                    "heading1": 1,
                    "heading2": 2,
                    "heading3": 3,
                    "heading4": 4,
                    "heading5": 5,
                    "heading6": 6,
                }

                for para in doc.paragraphs:
                    style_name = para.style.name.lower().strip() if para.style else ""
                    text = para.text.strip()

                    heading_level = None
                    for key, level in heading_style_map.items():
                        if key in style_name:
                            heading_level = level
                            break

                    if heading_level and text:
                        headings.append(Heading(level=heading_level, text=text))
                    elif text:
                        headings.append(Heading(level=7, text=text)) if self._is_heading_candidate(text) else paragraphs.append(
                            Paragraph(text=text)
                        )

                for table_index, table in enumerate(doc.tables, start=1):
                    rows: list[list[str]] = []
                    for row in table.rows:
                        rows.append([cell.text.strip() for cell in row.cells])
                    if rows:
                        headers = rows[0]
                        data_rows = rows[1:]
                        tables.append(
                            Table(
                                headers=headers,
                                rows=data_rows,
                                caption=f"Table {table_index}",
                            )
                        )

                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        images.append(
                            ImageMeta(
                                filename=rel.target_ref.split("/")[-1],
                                format=rel.target_ref.split(".")[-1] if "." in rel.target_ref else None,
                            )
                        )

        except Exception:
            pass

        if not title and headings:
            title = headings[0].text
        if not title and paragraphs:
            title = paragraphs[0].text[:120]
        if not title:
            title = filename.rsplit(".", 1)[0] if "." in filename else filename

        return ExtractedDocument(
            id=None,
            title=title,
            author=author if author else None,
            headings=headings,
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            code_blocks=code_blocks,
            metadata=metadata,
            source=filename,
            hash=doc_hash,
        )

    @staticmethod
    def _is_heading_candidate(text: str) -> bool:
        words = text.split()
        if len(words) > 12:
            return False
        if text.isupper() and len(text) < 100:
            return True
        if text.startswith(("Chapter", "CHAPTER")):
            return True
        return False
